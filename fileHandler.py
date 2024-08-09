from log import authLog
from docx import Document
from docx.shared import RGBColor
from auth import Auth
from commandsCLI import shCoreInfo, shIntDesSDW
from openpyxl import Workbook

import re
import os
import csv
import json
import traceback
import ipaddress
import openpyxl

removeCIDR_Patt = r'/\d{2}'
filterSiteCode = r'-sdw-0[1-9]'

PID_SDW03 = 'C8300-1N1S-4T2X-'
PID_SDW04 = 'C8300-1N1S-4T2X-'

ndlmPath1 = "NDLM_Template.xlsx"
ndlmPath2 = "NDLM_Tier2_Template.xlsx"

outputFolder = "Outputs"

sdw03Template = "sdw-03-template.csv"
sdw04Template = "sdw-04-template.csv"

returnList = []

def chooseCSV():
    csvDataList = []

    for i in range(2):
        while True:
            csvFile = input(f"Please enter the path to the CSV file {i + 1}: ")
            try:
                with open(csvFile, "r") as csvFileFinal:
                    authLog.info(f"User chose  the CSV File path: {csvFile}")
                    print(f"INFO: file successfully found.")
                    csvReader = csv.reader(csvFileFinal)
                    csvData = list(csvReader)
                    if csvData:
                        rowText = csvData[1]
                        for row in rowText:
                            print(f"{row}")
                        csvDataList.append(rowText)                         
                        break
                    else:
                        print(f"INFO: Cells not found under file: {csvFile}")
                        authLog.info(f"Cells not found under file: {csvFile}")
            except FileNotFoundError:
                print("File not found. Please check the file path and try again.")
                authLog.error(f"File not found in path {csvFile}")
                authLog.error(traceback.format_exc())
                continue

            except Exception as error:
                print(f"ERROR: {error}\n", traceback.format_exc())
                authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())
                
    mergedData = [item for sublist in csvDataList for item in sublist]
    for index, item in enumerate(mergedData):
        print(f"rowText[{index}] with string: {item}")
    os.system("PAUSE")
    return mergedData

def chooseDocx_ISR(rowText):
    swHostname, username, netDevice = Auth(rowText[12])
    shHostnameOut, netVlan1101, netVlan1103, shIntDesSDWOut, shIntDesCONOut1, shIntStatMPLSOut1, shVlanMgmtIP, shVlanMgmtCIDR, shLoop0Out = shCoreInfo(swHostname, username, netDevice)

    print(f"\n","="*76)
    print(f"INFO: Location: {rowText[3]}\n")

    print(f"INFO: BB1 Circuit Information: {rowText[65]}\n")

    print(f"INFO: MPLS Circuit Information {rowText[28]}")
    print(f"="*76, "\n")

    while True:
        try: 
            wordFile = "Caremore - Tier II - 8300 - vEdge to cEdge - gold.docx"
            wordDOC = Document(wordFile)
            authLog.info(f"User chose  the DOCX File path: {wordFile}")
            print(f"INFO: file successfully found: {wordFile}.")
            serialNumSDW01 = input("Please input the serial number of SDW-01: ")
            serialNumSDW02 = input("Please input the serial number of SDW-02: ")
            serialNumSDW03 = input("Please input the serial number of SDW-03: ")
            serialNumSDW04 = input("Please input the serial number of SDW-04: ")
            cEdge1Loop = input("Please input the SDW03 Loopback IP Address: ")
            cEdge2Loop = input("Please input the SDW04 Loopback IP Address: ")
            siteNo = input(f"Please input the new Site ID (Old Site ID: {rowText[41]}):")
            city = input("Please input the City: ")
            state = input("Please input the State: ")
            mplsCircuitID = input("Please input the MPLS Circuit ID:")
            bb1Carrier = input("Please input the bb1-carrier: ")
            bb1Circuitid = input("Please input the bb1-circuitid: ")
            cEdge2TLOC3_Port = input(f"Please input the cedge2-tloc3-port (TenGigabitEthernet0/0/5 or GigabitEthernet0/0/1 for {bb1Carrier} port): ")
            print("=" * 61,"\n\tINFO: Now begins information of the Core Switch")
            print("=" * 61)
            print(f"{shHostnameOut}{shIntDesSDW}\n{shIntDesSDWOut}\n")
            swcEdge1_vlan = input("Please input the VLAN for SDW-03, 1101 if possible: ")
            swcEdge2_vlan = input("Please input the VLAN for SDW-04, 1103 if possible: ")
            swcEdge1_port = input("Please input the connection to SDW-03 gi0/0/0 in VPN 1 from the switch: ")
            swcEdge2_port = input("Please input the connection to SDW-04 gi0/0/0 in VPN 1 from the switch: ")
            swcEdge1_mplsPort = input("Please input the Switch port for SDW-03 gi0/0/2 connection to Lumen: ")
            swcEdge2_mplsPort = input("Please input the Switch port for SDW-04 gi0/0/2 connection to Lumen: ")

            print("\nrowText 2:", rowText[2], "rowText 17:", rowText[17])
            print("After changes:")
            rowText[2] = re.sub('01', '03', rowText[2])
            rowText[17] = re.sub('02', '04', rowText[17])
            rowText[17] = re.sub('ge0/3', 'ge0/0/3', rowText[17])
            print("rowText 2:", rowText[2], "rowText 17:", rowText[17])
            os.system("PAUSE")

            print("\nrowText 44:", rowText[44], "rowText 59:", rowText[59])
            print("After changes:")
            rowText[44] = re.sub('02', '04', rowText[44])
            rowText[59] = re.sub('01', '03', rowText[59])
            rowText[59] = re.sub('ge0/3', 'ge0/0/3', rowText[59])
            print("rowText 44:", rowText[44], "rowText 59:", rowText[59])
            os.system("PAUSE")

            print("rowText 6:", rowText[6], "rowText 18:", rowText[18], "rowText 29:", rowText[29], \
                  f"rowText 48:", rowText[48], "rowText 79:", rowText[79],"")
            print("After changes:")
            rowText[6] = re.sub(removeCIDR_Patt, '', rowText[6])
            rowText[18] = re.sub(removeCIDR_Patt, '', rowText[18])
            rowText[29] = re.sub(removeCIDR_Patt, '', rowText[29])
            rowText[48] = re.sub(removeCIDR_Patt, '', rowText[48])
            rowText[79] = re.sub(removeCIDR_Patt, '', rowText[79])
            print("rowText 6:", rowText[6], "rowText 18:", rowText[18], "rowText 29:", rowText[29], \
                  f"rowText 48:", rowText[48], "rowText 79:", rowText[79],"")
            os.system("PAUSE")

            cedge2TLOC3_List = rowText[66]
            cedge2TLOC3_STR = ''.join(cedge2TLOC3_List)
            cedge2TLOC3_IP_STR = cedge2TLOC3_STR.split('/')[0]
            cedge2TLOC3_CIDR_STR = cedge2TLOC3_STR.split('/')[1]
            cedge2TLOC3_MASK_STR = ipaddress.IPv4Network(cedge2TLOC3_STR, strict=False)
            cedge2TLOC3_MASK_STR = str(cedge2TLOC3_MASK_STR.netmask)

            serialNumSDW03New = PID_SDW03 + serialNumSDW03
            serialNumSDW04New = PID_SDW04 + serialNumSDW04

            snmpLocation = f'{rowText[3]}'
            cedge1_host = f'{rowText[2]}'
            cedge2_host = f'{rowText[44]}'

            siteCode = f'{rowText[2]}'
            siteCode = re.sub(filterSiteCode, '', siteCode)
            print(f"This is the side code:{siteCode}")
            os.system("PAUSE")
            sw_host = f'{rowText[12]}'

            replaceText = {
                'cedge1-host' : f'{rowText[2]}',
                'snmp-location' : f'{rowText[3]}',
                'cedge1-rtr-ip' : f'{rowText[6]}',
                'cEdge-asn' : f'{rowText[8]}',
                'cedge1-sw-ip' : f'{rowText[11]}',
                'switch-asn' : f'{rowText[13]}',
                'mpls-pe-ip' : f'{rowText[14]}',
                'cedge2-tloc3-ext-ip' : f'{rowText[15]}',
                'cedge2-host - gi0/0/3 - TLOC3' : f'{rowText[17]}',
                'cedge1-tloc3-ip'	: f'{rowText[18]}',
                'mpls-ce1-ip' : f'{rowText[29]}',
                'mpls-speed' : f'{rowText[35]}',
                'latitude' : f'{rowText[38]}',
                'longitude' : f'{rowText[39]}',
                # Here starts the second CSV file #
                'cedge2-host'	: f'{rowText[44]}',
                'bb1-down-speed' : f'{rowText[76]}',
                'cedge2-rtr-ip' : f'{rowText[48]}',
                'cedge2-sw-ip' : f'{rowText[53]}',	
                'cedge2-tloc3-gate' : f'{rowText[57]}',	
                'cedge1-host TLOC3 gi0/0/3' : f'{rowText[59]}',
                'cedge2-tloc3-ext-ip/30' : f'{rowText[60]}',
                'bb1-up-speed' : f'{rowText[75]}',	
                'mpls-ce2-ip'	: f'{rowText[79]}'
            }

            print(json.dumps(replaceText, indent=4))
            os.system("PAUSE")

            stringRegexPatt = {
                'cedge1-serial-no' : serialNumSDW03New,
                'cedge2-serial-no' : serialNumSDW04New,
                'cedge1-loop' : cEdge1Loop,
                'cedge2-loop' : cEdge2Loop,
                'site-no'	: siteNo,
                'city': city,
                'state': state,
                'site-code': siteCode,
                'sw-mgmt-ip' : shVlanMgmtIP,
                'sw-host' : sw_host,
                'sw-cEdge1-mpls-port': swcEdge1_mplsPort,
                'sw-cEdge2-mpls-port': swcEdge2_mplsPort,
                'mpls-circuitid':  mplsCircuitID,
                'bb1-carrier': bb1Carrier,
                'bb1-circuitid': bb1Circuitid,
                'cedge2-tloc3-port': cEdge2TLOC3_Port,
                'cedge2-tloc3-ip': cedge2TLOC3_IP_STR,
                'cedge2-tloc3-mask' : cedge2TLOC3_MASK_STR,
                'cedge2-tloc3-cidr': cedge2TLOC3_CIDR_STR,
                'cedge1-lan-net': netVlan1101,
                'cedge2-lan-net': netVlan1103,
                'sw-loop': shLoop0Out,
                'sw-mgmt-cidr': shVlanMgmtCIDR,
                'sw-cedge1-port': swcEdge1_port,
                'sw-cedge1-vlan': swcEdge1_vlan,
                'sw-cedge2-port': swcEdge2_port,
                'sw-cedge2-vlan': swcEdge2_vlan,
                'sw-mpls-port': shIntStatMPLSOut1[0],
                'sw-remote-con-net1': shIntDesCONOut1[0],
                'sw-remote-con-net2': shIntDesCONOut1[1],
                'sw-mgmt-vlan' : '1500'
            }

            manualReplacements = {re.compile(r'\b{}\b'.format(pattern), re.IGNORECASE): value for pattern, value in stringRegexPatt.items()}

            for para in wordDOC.paragraphs:
                if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in para.runs):
                    print(f"Found red text: {para.text}")
                    for wordString, csvString in replaceText.items():
                        if re.search(r'\b{}\b'.format(re.escape(wordString)), para.text, re.IGNORECASE):
                            print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                            authLog.info(f"Replacing '{wordString}' with '{csvString}'")
                            para.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, para.text, flags=re.IGNORECASE)

                    for placeholder, replacement in manualReplacements.items():
                        replacement = str(replacement)
                        if placeholder.search(para.text):
                            print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            authLog.info(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            para.text = placeholder.sub(replacement, para.text)

            for table in wordDOC.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in paragraph.runs):
                                print(f"Found red text: {paragraph.text}")
                                for wordString, csvString in replaceText.items():
                                    if re.search(r'\b{}\b'.format(re.escape(wordString)), paragraph.text, re.IGNORECASE):
                                        print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                                        authLog.info(f"Replacing in Table: '{wordString}' with '{csvString}'")
                                        paragraph.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, paragraph.text, flags=re.IGNORECASE)

                                for placeholder, replacement in manualReplacements.items():
                                    replacement = str(replacement)
                                    if placeholder.search(paragraph.text):
                                        print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                                        authLog.info(f"Replacing in Table: '{placeholder.pattern}' with '{replacement}'")
                                        paragraph.text = placeholder.sub(replacement, paragraph.text)

            newWordDoc = f"Outputs/{siteCode}_ImplementationPlan.docx"
            wordDOC.save(newWordDoc)
            authLog.info(f"Replacements made successfully in DOCX file and saved as: {newWordDoc}")
            print(f"INFO: Replacements made successfully in DOCX file and saved as: {newWordDoc}")

            manualReplaceList = [
                serialNumSDW01,     #0
                serialNumSDW02,     #1
                serialNumSDW03,     #2
                serialNumSDW04,     #3
                cEdge1Loop,         #4
                cEdge2Loop,         #5
                siteNo,             #6
                city,               #7
                state,              #8
                siteCode,           #9
                shVlanMgmtIP,       #10            
                swcEdge1_mplsPort,  #11
                swcEdge2_mplsPort,  #12
                mplsCircuitID,      #13
                bb1Carrier,         #14
                bb1Circuitid,       #15
                cEdge2TLOC3_Port,   #16
                cedge2TLOC3_IP_STR, #17
                cedge2TLOC3_MASK_STR,#18
                cedge2TLOC3_CIDR_STR,#19
                netVlan1101,        #20
                netVlan1103,        #21
                shLoop0Out,         #22
                shVlanMgmtCIDR,     #23
                swcEdge1_port,      #24
                swcEdge1_vlan,      #25
                swcEdge2_port,      #26
                swcEdge2_vlan,      #27
                shIntStatMPLSOut1[0],#28
                shIntDesCONOut1[0], #29
                shIntDesCONOut1[1], #30
                sw_host,            #31
                '1500'              #32
            ]

            return {
                'rowText' : rowText,
                'rowText1' :  manualReplaceList
            }

        except FileNotFoundError:
            print("File not found. Please check the file path and try again.")
            authLog.error(f"File not found in path {wordFile}")
            authLog.error(traceback.format_exc())
            continue

        except Exception as error:
            print(f"ERROR: {error}\n", traceback.format_exc())
            authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def modNDLMISR(rowText, rowText1):
    try:
        replaceText = {
            'site-code' : f'{rowText1[9]}',
            'vedge1-serial-no' : f'{rowText1[0]}',
            'vedge2-serial-no' : f'{rowText1[1]}',
            'cedge1-serial-no' : f'{rowText1[2]}',
            'cedge2-serial-no' : f'{rowText1[3]}',
            'cedge1-loop' : f'{rowText1[4]}',
            'cedge2-loop' : f'{rowText1[5]}',
            'snmp-location' : f'{rowText[3]}',
            'vedge1-loop': f'{rowText[1]}',
            'vedge2-loop': f'{rowText[43]}'
        }

        ndlmFile = openpyxl.load_workbook(ndlmPath1)
        ndlmFileSheet = ndlmFile.active

        for row in ndlmFileSheet.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    cell.value = cellValue

            newNDLMFile = os.path.join(outputFolder, f'{rowText1[9]}-NDLM.xlsx')
            ndlmFile.save(newNDLMFile)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath1}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def modNDLM2ISR(rowText, rowText1):
    try:

        replaceText = {
            'site-code' : f'{rowText1[9]}',
            'cedge1-loop' : f'{rowText1[4]}',
            'cedge2-loop' : f'{rowText1[5]}',
            'snmp-location' : f'{rowText[3]}',
            'city': f'{rowText1[7]}',
            'state': f'{rowText1[8]}',
            'site-no': f'{rowText1[6]}',
            'cedge1-host': f'{rowText[2]}',
            'cedge2-host': f'{rowText[44]}',
            'sw-host' : f'{rowText1[31]}',
            'sw-mpls-port' : f'{rowText1[28]}',
            'cedge2-tloc3-port': f'{rowText1[16]}',
            'sw-cedge1-port' : f'{rowText1[24]}',
            'sw-cedge2-port' : f'{rowText1[26]}',
            'sw-cedge1-mpls-port' : f'{rowText1[11]}',
            'sw-cedge2-mpls-port' : f'{rowText1[12]}'
        }

        ndlmFile1 = openpyxl.load_workbook(ndlmPath2)
        ndlmFileSheet1 = ndlmFile1.active

        for row in ndlmFileSheet1.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    cell.value = cellValue

            newNDLMFile1 = os.path.join(outputFolder, f'{rowText1[9]}-NDLM-Tier2.xlsx')
            ndlmFile1.save(newNDLMFile1)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath2}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def cEdgeTemplateISR(rowText, rowText1):

    for index, item in enumerate(rowText):
        print(f"rowText[{index}] with string: {item}")
    
    for index, item in enumerate(rowText1):
        print(f"rowText1[{index}] with string: {item}")
    os.system("PAUSE")
    
    newSDW03Template = f'Outputs/{rowText1[9]}-SDW-03-Template.csv'
    newSDW04Template = f'Outputs/{rowText1[9]}-SDW-04-Template.csv'

    sdw03Replacements = {
        'cedge1-host' : f'{rowText[2]}',
        'snmp-location' : f'{rowText[3]}',
        'cedge1-rtr-ip' : f'{rowText[6]}',
        'cEdge-asn' : f'{rowText[8]}',
        'cedge1-sw-ip' : f'{rowText[11]}',
        'switch-asn' : f'{rowText[13]}',
        'mpls-pe-ip' : f'{rowText[14]}',
        'cedge2-tloc3-ext-ip' : f'{rowText[15]}',
        'cedge2-host - gi0/0/3 - TLOC3' : f'{rowText[17]}',
        'cedge1-tloc3-ip'	: f'{rowText[18]}',
        'mpls-ce1-ip' : f'{rowText[29]}',
        'mpls-speed' : f'{rowText[35]}',
        'latitude' : f'{rowText[38]}',
        'longitude' : f'{rowText[39]}',
        # Here starts the second CSV file #
        'cedge2-host'	: f'{rowText[44]}',
        'bb1-down-speed' : f'{rowText[76]}',
        'cedge2-rtr-ip' : f'{rowText[48]}',
        'cedge2-sw-ip' : f'{rowText[53]}',	
        'cedge2-tloc3-gate' : f'{rowText[57]}',	
        'cedge1-host TLOC3 gi0/0/3' : f'{rowText[59]}',
        'cedge2-tloc3-ext-ip/30' : f'{rowText[60]}',
        'bb1-up-speed' : f'{rowText[75]}',	
        'mpls-ce2-ip'	: f'{rowText[79]}',

        'cedge1-serial-no' : f'{rowText1[2]}',
        'cedge2-serial-no' : f'{rowText1[3]}',
        'cedge1-loop' : f'{rowText1[4]}',
        'cedge2-loop' : f'{rowText1[5]}',
        'site-no'	: f'{rowText1[6]}',
        'city': f'{rowText1[7]}',
        'state': f'{rowText1[8]}',
        'site-code': f'{rowText1[9]}',
        'sw-mgmt-ip' : f'{rowText1[10]}',
        'sw-host' : f'{rowText1[31]}',
        'sw-cEdge1-mpls-port': f'{rowText1[11]}',
        'sw-cEdge2-mpls-port': f'{rowText1[12]}',
        'mpls-circuitid':  f'{rowText1[13]}',
        'bb1-carrier': f'{rowText1[14]}',
        'bb1-circuitid': f'{rowText1[15]}',
        'cedge2-tloc3-port': f'{rowText1[16]}',
        'cedge2-tloc3-ip': f'{rowText1[17]}',
        'cedge2-tloc3-mask' : f'{rowText1[18]}',
        'cedge2-tloc3-cidr': f'{rowText1[19]}',
        'cedge1-lan-net': f'{rowText1[20]}',
        'cedge2-lan-net': f'{rowText1[21]}',
        'sw-loop': f'{rowText1[22]}',
        'sw-mgmt-cidr': f'{rowText1[23]}',
        'sw-cedge1-port': f'{rowText1[24]}',
        'sw-cedge1-vlan': f'{rowText1[25]}',
        'sw-cedge2-port': f'{rowText1[26]}',
        'sw-cedge2-vlan': f'{rowText1[27]}',
        'sw-mpls-port': f'{rowText1[28]}',
        'sw-remote-con-net1': f'{rowText1[29]}',
        'sw-remote-con-net2': f'{rowText1[30]}',
        'sw-mgmt-vlan' : f'{rowText1[32]}'        
    }

    sdw04Replacements = {
 
    }

    try:
        with open(sdw03Template, "r") as inputCSV:
            authLog.info(f"Generating {rowText1[9]}-SDW-03-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-03-Template.")
            csvReader = csv.reader(inputCSV)
               
            rows = list(csvReader)

            if len(rows) > 1:
                secondRow = rows[1]
                modifiedRow = []
                for cell in secondRow:
                    cellValue = str(cell).strip()
                    for key, value in sdw03Replacements.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    modifiedRow.append(cellValue)
                rows[1] = modifiedRow

        with open(newSDW03Template, 'w', newline="") as outputCSV:
            csvWriter = csv.writer(outputCSV)
            csvWriter.writerows(rows)
    
        with open(sdw04Template, "r") as inputCSV1:
            authLog.info(f"Generating {rowText1[9]}-SDW-04-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-04-Template.")
            csvReader1 = csv.reader(inputCSV1)
               
            rows1 = list(csvReader1)

            if len(rows1) > 1:
                secondRow1 = rows1[1]
                modifiedRow1 = []
                for cell1 in secondRow1:
                    cellValue1 = str(cell1).strip()
                    for key1, value1 in sdw04Replacements.items():
                        if key1.lower() in cellValue1.lower():
                            cellValue1 = cellValue1.replace(key1, value1)
                    modifiedRow1.append(cellValue1)
                rows1[1] = modifiedRow1
            
        with open(newSDW04Template, 'w', newline="") as outputCSV1:
            csvWriter1 = csv.writer(outputCSV1)
            csvWriter1.writerows(rows1)

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Error message: {error}\n", traceback.format_exc())

def chooseDocx_vEdge(rowText):
    swHostname, username, netDevice = Auth(rowText[13])
    shHostnameOut, netVlan1101, netVlan1103, shIntDesSDWOut, shIntDesCONOut1, shIntStatMPLSOut1, shVlanMgmtIP, shVlanMgmtCIDR, shLoop0Out = shCoreInfo(swHostname, username, netDevice)

    print(f"\n","="*76)
    print(f"INFO: Location: {rowText[3]}\n")

    print(f"INFO: BB1 Circuit Information: {rowText[71]}\n")

    print(f"INFO: MPLS Circuit Information {rowText[31]}")
    print(f"="*76, "\n")

    while True:
        try:
            wordFile = "Caremore - Tier II - 8300 - vEdge to cEdge - gold.docx"
            wordDOC = Document(wordFile)
            authLog.info(f"User chose  the DOCX File path: {wordFile}")
            print(f"INFO: file successfully found: {wordFile}.")
            serialNumSDW01 = input("Please input the serial number of SDW-01: ")
            serialNumSDW02 = input("Please input the serial number of SDW-02: ")
            serialNumSDW03 = input("Please input the serial number of SDW-03: ")
            serialNumSDW04 = input("Please input the serial number of SDW-04: ")
            cEdge1Loop = input("Please input the SDW03 Loopback IP Address: ")
            cEdge2Loop = input("Please input the SDW04 Loopback IP Address: ")
            siteNo = input(f"Please input the new Site ID (Old Site ID: {rowText[44]}):")
            city = input("Please input the City: ")
            state = input("Please input the State: ")
            mplsCircuitID = input("Please input the MPLS Circuit ID:")
            bb1Carrier = input("Please input the bb1-carrier: ")
            bb1Circuitid = input("Please input the bb1-circuitid: ")
            cEdge2TLOC3_Port = input(f"Please input the cedge2-tloc3-port (TenGigabitEthernet0/0/5 or GigabitEthernet0/0/1 for {bb1Carrier} port): ")
            print("=" * 61,"\n\tINFO: Now begins information of the Core Switch")
            print("=" * 61)
            print(f"{shHostnameOut}{shIntDesSDW}\n{shIntDesSDWOut}\n")
            swcEdge1_vlan = input("Please input the VLAN for SDW-03, 1101 if possible: ")
            swcEdge2_vlan = input("Please input the VLAN for SDW-04, 1103 if possible: ")
            swcEdge1_port = input("Please input the connection to SDW-03 gi0/0/0 in VPN 1 from the switch: ")
            swcEdge2_port = input("Please input the connection to SDW-04 gi0/0/0 in VPN 1 from the switch: ")
            swcEdge1_mplsPort = input("Please input the Switch port for SDW-03 gi0/0/2 connection to Lumen: ")
            swcEdge2_mplsPort = input("Please input the Switch port for SDW-04 gi0/0/2 connection to Lumen: ")

            print("\nrowText 2:", rowText[2], "rowText 20:", rowText[20])
            print("After changes:")
            rowText[2] = re.sub('01', '03', rowText[2])
            rowText[20] = re.sub('02', '04', rowText[20])
            rowText[20] = re.sub('ge0/3', 'ge0/0/3', rowText[20])
            print("rowText 2:", rowText[2], "rowText 20:", rowText[20])
            os.system("PAUSE")

            print("\nrowText 47:", rowText[47], "rowText 65:", rowText[65])
            print("After changes:")
            rowText[47] = re.sub('02', '04', rowText[47])
            rowText[65] = re.sub('01', '03', rowText[65])
            rowText[65] = re.sub('ge0/3', 'ge0/0/3', rowText[65])
            print("rowText 47:", rowText[47], "rowText 65:", rowText[65])
            os.system("PAUSE")

            print("rowText 6:", rowText[6], "rowText 21:", rowText[21], "rowText 32:", rowText[32], \
                  f"rowText 51:", rowText[51], "rowText 85:", rowText[85],"")
            print("After changes:")
            rowText[6] = re.sub(removeCIDR_Patt, '', rowText[6])
            rowText[21] = re.sub(removeCIDR_Patt, '', rowText[21])
            rowText[32] = re.sub(removeCIDR_Patt, '', rowText[32])
            rowText[51] = re.sub(removeCIDR_Patt, '', rowText[51])
            rowText[85] = re.sub(removeCIDR_Patt, '', rowText[85])
            print("rowText 6:", rowText[6], "rowText 21:", rowText[21], "rowText 32:", rowText[32], \
                  f"rowText 51:", rowText[51], "rowText 85:", rowText[85],"")
            os.system("PAUSE")

            cedge2TLOC3_List = rowText[72]
            cedge2TLOC3_STR = ''.join(cedge2TLOC3_List)
            cedge2TLOC3_IP_STR = cedge2TLOC3_STR.split('/')[0]
            cedge2TLOC3_CIDR_STR = cedge2TLOC3_STR.split('/')[1]
            cedge2TLOC3_MASK_STR = ipaddress.IPv4Network(cedge2TLOC3_STR, strict=False)
            cedge2TLOC3_MASK_STR = str(cedge2TLOC3_MASK_STR.netmask)

            serialNumSDW03New = PID_SDW03 + serialNumSDW03
            serialNumSDW04New = PID_SDW04 + serialNumSDW04

            siteCode = f'{rowText[2]}'
            siteCode = re.sub(filterSiteCode, '', siteCode)
            print(f"This is the side code:{siteCode}")
            os.system("PAUSE")
            sw_host = f'{rowText[13]}'

            replaceText = {
                'cedge1-host' : f'{rowText[2]}',
                'snmp-location' : f'{rowText[3]}',
                'cedge1-rtr-ip' : f'{rowText[6]}',
                'cEdge-asn' : f'{rowText[9]}',
                'cedge1-sw-ip' : f'{rowText[12]}',
                'switch-asn' : f'{rowText[14]}',
                'mpls-pe-ip' : f'{rowText[17]}',
                'cedge2-tloc3-ext-ip' : f'{rowText[18]}',
                'cedge2-host - gi0/0/3 - TLOC3' : f'{rowText[20]}',
                'cedge1-tloc3-ip'	: f'{rowText[21]}',
                'mpls-ce1-ip' : f'{rowText[32]}',
                'mpls-speed' : f'{rowText[38]}',
                'latitude' : f'{rowText[41]}',
                'longitude' : f'{rowText[42]}',
                # Here starts the second CSV file #
                'cedge2-host'	: f'{rowText[47]}',
                'bb1-down-speed' : f'{rowText[82]}',
                'cedge2-rtr-ip' : f'{rowText[51]}',
                'cedge2-sw-ip' : f'{rowText[57]}',	
                'cedge2-tloc3-gate' : f'{rowText[63]}',	
                'cedge1-host TLOC3 gi0/0/3' : f'{rowText[59]}',
                'cedge2-tloc3-ext-ip/30' : f'{rowText[60]}',
                'bb1-up-speed' : f'{rowText[82]}',	
                'mpls-ce2-ip'	: f'{rowText[85]}'
            }

            print(json.dumps(replaceText, indent=4))
            os.system("PAUSE")

            stringRegexPatt = {
                'cedge1-serial-no' : serialNumSDW03New,
                'cedge2-serial-no' : serialNumSDW04New,
                'cedge1-loop' : cEdge1Loop,
                'cedge2-loop' : cEdge2Loop,
                'site-no'	: siteNo,
                'city': city,
                'state': state,
                'site-code': siteCode,
                'sw-mgmt-ip' : shVlanMgmtIP,
                'sw-cEdge1-mpls-port': swcEdge1_mplsPort,
                'sw-cEdge2-mpls-port': swcEdge2_mplsPort,
                'mpls-circuitid':  mplsCircuitID,
                'bb1-carrier': bb1Carrier,
                'bb1-circuitid': bb1Circuitid,
                'cedge2-tloc3-port': cEdge2TLOC3_Port,
                'cedge2-tloc3-ip': cedge2TLOC3_IP_STR,
                'cedge2-tloc3-mask' : cedge2TLOC3_MASK_STR,
                'cedge2-tloc3-cidr': cedge2TLOC3_CIDR_STR,
                'cedge1-lan-net': netVlan1101,
                'cedge2-lan-net': netVlan1103,
                'sw-loop': shLoop0Out,
                'sw-mgmt-cidr': shVlanMgmtCIDR,
                'sw-cedge1-port': swcEdge1_port,
                'sw-cedge1-vlan': swcEdge1_vlan,
                'sw-cedge2-port': swcEdge2_port,
                'sw-cedge2-vlan': swcEdge2_vlan,
                'sw-mpls-port': shIntStatMPLSOut1[0],
                'sw-remote-con-net1': shIntDesCONOut1[0],
                'sw-remote-con-net2': shIntDesCONOut1[1],
                'sw-host' : sw_host,
                'sw-mgmt-vlan' : '1500'
            }

            manualReplacements = {re.compile(r'\b{}\b'.format(pattern), re.IGNORECASE): value for pattern, value in stringRegexPatt.items()}

            for para in wordDOC.paragraphs:
                if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in para.runs):
                    print(f"Found red text: {para.text}")
                    for wordString, csvString in replaceText.items():
                        if re.search(r'\b{}\b'.format(re.escape(wordString)), para.text, re.IGNORECASE):
                            print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                            authLog.info(f"Replacing '{wordString}' with '{csvString}'")
                            para.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, para.text, flags=re.IGNORECASE)

                    for placeholder, replacement in manualReplacements.items():
                        replacement = str(replacement)
                        if placeholder.search(para.text):
                            print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            authLog.info(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                            para.text = placeholder.sub(replacement, para.text)

            for table in wordDOC.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if any(run.font.color.rgb == RGBColor(255, 0, 0) for run in paragraph.runs):
                                print(f"Found red text: {paragraph.text}")
                                for wordString, csvString in replaceText.items():
                                    if re.search(r'\b{}\b'.format(re.escape(wordString)), paragraph.text, re.IGNORECASE):
                                        print(f"INFO: Replacing '{wordString}' with '{csvString}'")
                                        authLog.info(f"Replacing in Table: '{wordString}' with '{csvString}'")
                                        paragraph.text = re.sub(r'\b{}\b'.format(re.escape(wordString)), csvString, paragraph.text, flags=re.IGNORECASE)

                                for placeholder, replacement in manualReplacements.items():
                                    replacement = str(replacement)
                                    if placeholder.search(paragraph.text):
                                        print(f"Replacing '{placeholder.pattern}' with '{replacement}'")
                                        authLog.info(f"Replacing in Table: '{placeholder.pattern}' with '{replacement}'")
                                        paragraph.text = placeholder.sub(replacement, paragraph.text)

            newWordDoc = f"Outputs/{siteCode}_ImplementationPlan.docx"
            wordDOC.save(newWordDoc)
            authLog.info(f"Replacements made successfully in DOCX file and saved as: {newWordDoc}")
            print(f"INFO: Replacements made successfully in DOCX file and saved as: {newWordDoc}")
            
            os.system("PAUSE")

            manualReplaceList = [
                serialNumSDW01,     #0
                serialNumSDW02,     #1
                serialNumSDW03,     #2
                serialNumSDW04,     #3
                cEdge1Loop,         #4
                cEdge2Loop,         #5
                siteNo,             #6
                city,               #7
                state,              #8
                siteCode,           #9
                shVlanMgmtIP,       #10            
                swcEdge1_mplsPort,  #11
                swcEdge2_mplsPort,  #12
                mplsCircuitID,      #13
                bb1Carrier,         #14
                bb1Circuitid,       #15
                cEdge2TLOC3_Port,   #16
                cedge2TLOC3_IP_STR, #17
                cedge2TLOC3_MASK_STR,#18
                cedge2TLOC3_CIDR_STR,#19
                netVlan1101,        #20
                netVlan1103,        #21
                shLoop0Out,         #22
                shVlanMgmtCIDR,     #23
                swcEdge1_port,      #24
                swcEdge1_vlan,      #25
                swcEdge2_port,      #26
                swcEdge2_vlan,      #27
                shIntStatMPLSOut1[0],#28
                shIntDesCONOut1[0], #29
                shIntDesCONOut1[1], #30
                sw_host,            #31
                '1500'              #32

            ]

            return {
                'rowText' : rowText,
                'rowText1' :  manualReplaceList
            }

        
        except FileNotFoundError:
            print("File not found. Please check the file path and try again.")
            authLog.error(f"File not found in path {wordFile}")
            authLog.error(traceback.format_exc())
            continue

        except Exception as error:
            print(f"ERROR: {error}\n", traceback.format_exc())
            authLog.error(f"Wasn't possible to choose the DOCX file, error message: {error}\n{traceback.format_exc()}")

def modNDLMvEdge(rowText, rowText1):
    try:
        replaceText = {
            'site-code' : f'{rowText1[9]}',
            'vedge1-serial-no' : f'{rowText1[0]}',
            'vedge2-serial-no' : f'{rowText1[1]}',
            'cedge1-serial-no' : f'{rowText1[2]}',
            'cedge2-serial-no' : f'{rowText1[3]}',
            'cedge1-loop' : f'{rowText1[4]}',
            'cedge2-loop' : f'{rowText1[5]}',
            'snmp-location' : f'{rowText[3]}',
            'vedge1-loop': f'{rowText[1]}',
            'vedge2-loop': f'{rowText[46]}'
        }

        ndlmFile = openpyxl.load_workbook(ndlmPath1)
        ndlmFileSheet = ndlmFile.active

        for row in ndlmFileSheet.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    cell.value = cellValue

            newNDLMFile = os.path.join(outputFolder, f'{rowText1[9]}-NDLM.xlsx')
            ndlmFile.save(newNDLMFile)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath1}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def modNDLM2vEdge(rowText, rowText1):
    try:

        replaceText = {
            'site-code' : f'{rowText1[9]}',
            'cedge1-loop' : f'{rowText1[4]}',
            'cedge2-loop' : f'{rowText1[5]}',
            'snmp-location' : f'{rowText[3]}',
            'city': f'{rowText1[7]}',
            'state': f'{rowText1[8]}',
            'site-no': f'{rowText1[6]}',
            'cedge1-host': f'{rowText[2]}',
            'cedge2-host': f'{rowText[47]}',
            'sw-host' : f'{rowText1[31]}',
            'sw-mpls-port' : f'{rowText1[28]}',
            'cedge2-tloc3-port': f'{rowText1[16]}',
            'sw-cedge1-port' : f'{rowText1[24]}',
            'sw-cedge2-port' : f'{rowText1[26]}',
            'sw-cedge1-mpls-port' : f'{rowText1[11]}',
            'sw-cedge2-mpls-port' : f'{rowText1[12]}'
        }

        ndlmFile1 = openpyxl.load_workbook(ndlmPath2)
        ndlmFileSheet1 = ndlmFile1.active

        for row in ndlmFileSheet1.iter_rows():
            for cell in row:
                if cell.value:
                    cellValue = str(cell.value).strip()
                    for key, value in replaceText.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    cell.value = cellValue

            newNDLMFile1 = os.path.join(outputFolder, f'{rowText1[9]}-NDLM-Tier2.xlsx')
            ndlmFile1.save(newNDLMFile1)

    except FileNotFoundError:
        print("File not found. Please check the file path and try again.")
        authLog.error(f"File not found in path {ndlmPath2}")
        authLog.error(traceback.format_exc())

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Wasn't possible to choose the CSV file, error message: {error}\n", traceback.format_exc())

def cEdgeTemplatevEdge(rowText, rowText1):

    for index, item in enumerate(rowText):
        print(f"rowText[{index}] with string: {item}")
    
    for index, item in enumerate(rowText1):
        print(f"rowText1[{index}] with string: {item}")
    os.system("PAUSE")
    
    newSDW03Template = f'Outputs/{rowText1[9]}-SDW-03-Template.csv'
    newSDW04Template = f'Outputs/{rowText1[9]}-SDW-04-Template.csv'

    sdw03Replacements = {
        'cedge1-host' : f'{rowText[2]}',
        'snmp-location' : f'{rowText[3]}',
        'cedge1-rtr-ip' : f'{rowText[6]}',
        'cEdge-asn' : f'{rowText[9]}',
        'cedge1-sw-ip' : f'{rowText[12]}',
        'switch-asn' : f'{rowText[14]}',
        'mpls-pe-ip' : f'{rowText[17]}',
        'cedge2-tloc3-ext-ip' : f'{rowText[18]}',
        'cedge2-host - gi0/0/3 - TLOC3' : f'{rowText[20]}',
        'cedge1-tloc3-ip'	: f'{rowText[21]}',
        'mpls-ce1-ip' : f'{rowText[32]}',
        'mpls-speed' : f'{rowText[38]}',
        'latitude' : f'{rowText[41]}',
        'longitude' : f'{rowText[42]}',
        # Here starts the second CSV file #
        'cedge2-host'	: f'{rowText[47]}',
        'bb1-down-speed' : f'{rowText[82]}',
        'cedge2-rtr-ip' : f'{rowText[51]}',
        'cedge2-sw-ip' : f'{rowText[57]}',	
        'cedge2-tloc3-gate' : f'{rowText[63]}',	
        'cedge1-host TLOC3 gi0/0/3' : f'{rowText[59]}',
        'cedge2-tloc3-ext-ip/30' : f'{rowText[60]}',
        'bb1-up-speed' : f'{rowText[82]}',	
        'mpls-ce2-ip'	: f'{rowText[85]}',

        'cedge1-serial-no' : rowText1[2],
        'cedge2-serial-no' : rowText1[3],
        'cedge1-loop' : rowText1[4],
        'cedge2-loop' : rowText1[5],
        'site-no'	: rowText1[6],
        'city': rowText1[7],
        'state': rowText1[8],
        'site-code': rowText1[9],
        'sw-mgmt-ip' : rowText1[10],
        'sw-cEdge1-mpls-port': rowText[11],
        'sw-cEdge2-mpls-port': rowText1[12],
        'mpls-circuitid':  rowText1[13],
        'bb1-carrier': rowText1[14],
        'bb1-circuitid': rowText1[15],
        'cedge2-tloc3-port': rowText1[16],
        'cedge2-tloc3-ip': rowText1[17],
        'cedge2-tloc3-mask' : rowText1[18],
        'cedge2-tloc3-cidr': rowText1[19],
        'cedge1-lan-net': rowText1[20],
        'cedge2-lan-net': rowText1[21],
        'sw-loop': rowText1[22],
        'sw-mgmt-cidr': rowText1[23],
        'sw-cedge1-port': rowText1[24],
        'sw-cedge1-vlan': rowText1[25],
        'sw-cedge2-port': rowText1[26],
        'sw-cedge2-vlan': rowText1[27],
        'sw-mpls-port': rowText1[28],
        'sw-remote-con-net1': rowText1[29],
        'sw-remote-con-net2': rowText1[30],
        'sw-host' : rowText1[31],
        'sw-mgmt-vlan' : rowText1[32] 
    }

    sdw04Replacements = {
        'cedge1-host' : f'{rowText[2]}',
        'snmp-location' : f'{rowText[3]}',
        'cedge1-rtr-ip' : f'{rowText[6]}',
        'cEdge-asn' : f'{rowText[9]}',
        'cedge1-sw-ip' : f'{rowText[12]}',
        'switch-asn' : f'{rowText[14]}',
        'mpls-pe-ip' : f'{rowText[17]}',
        'cedge2-tloc3-ext-ip' : f'{rowText[18]}',
        'cedge2-host - gi0/0/3 - TLOC3' : f'{rowText[20]}',
        'cedge1-tloc3-ip'	: f'{rowText[21]}',
        'mpls-ce1-ip' : f'{rowText[32]}',
        'mpls-speed' : f'{rowText[38]}',
        'latitude' : f'{rowText[41]}',
        'longitude' : f'{rowText[42]}',
        # Here starts the second CSV file #
        'cedge2-host'	: f'{rowText[47]}',
        'bb1-down-speed' : f'{rowText[82]}',
        'cedge2-rtr-ip' : f'{rowText[51]}',
        'cedge2-sw-ip' : f'{rowText[57]}',	
        'cedge2-tloc3-gate' : f'{rowText[63]}',	
        'cedge1-host TLOC3 gi0/0/3' : f'{rowText[59]}',
        'cedge2-tloc3-ext-ip/30' : f'{rowText[60]}',
        'bb1-up-speed' : f'{rowText[82]}',	
        'mpls-ce2-ip'	: f'{rowText[85]}',
        
        'cedge1-serial-no' : rowText1[2],
        'cedge2-serial-no' : rowText1[3],
        'cedge1-loop' : rowText1[4],
        'cedge2-loop' : rowText1[5],
        'site-no'	: rowText1[6],
        'city': rowText1[7],
        'state': rowText1[8],
        'site-code': rowText1[9],
        'sw-mgmt-ip' : rowText1[10],
        'sw-cEdge1-mpls-port': rowText[11],
        'sw-cEdge2-mpls-port': rowText1[12],
        'mpls-circuitid':  rowText1[13],
        'bb1-carrier': rowText1[14],
        'bb1-circuitid': rowText1[15],
        'cedge2-tloc3-port': rowText1[16],
        'cedge2-tloc3-ip': rowText1[17],
        'cedge2-tloc3-mask' : rowText1[18],
        'cedge2-tloc3-cidr': rowText1[19],
        'cedge1-lan-net': rowText1[20],
        'cedge2-lan-net': rowText1[21],
        'sw-loop': rowText1[22],
        'sw-mgmt-cidr': rowText1[23],
        'sw-cedge1-port': rowText1[24],
        'sw-cedge1-vlan': rowText1[25],
        'sw-cedge2-port': rowText1[26],
        'sw-cedge2-vlan': rowText1[27],
        'sw-mpls-port': rowText1[28],
        'sw-remote-con-net1': rowText1[29],
        'sw-remote-con-net2': rowText1[30],
        'sw-host' : rowText1[31],
        'sw-mgmt-vlan' : rowText1[32] 
    }

    try:
        with open(sdw03Template, "r") as inputCSV:
            authLog.info(f"Generating {rowText1[9]}-SDW-03-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-03-Template.")
            csvReader = csv.reader(inputCSV)
               
            rows = list(csvReader)

            if len(rows) > 1:
                secondRow = rows[1]
                modifiedRow = []
                for cell in secondRow:
                    cellValue = str(cell).strip()
                    for key, value in sdw03Replacements.items():
                        if key.lower() in cellValue.lower():
                            cellValue = cellValue.replace(key, value)
                    modifiedRow.append(cellValue)
                rows[1] = modifiedRow

        with open(newSDW03Template, 'w', newline="") as outputCSV:
            csvWriter = csv.writer(outputCSV)
            csvWriter.writerows(rows)
    
        with open(sdw04Template, "r") as inputCSV1:
            authLog.info(f"Generating {rowText1[9]}-SDW-04-Template")
            print(f"INFO: Generating {rowText1[9]}-SDW-04-Template.")
            csvReader1 = csv.reader(inputCSV1)
               
            rows1 = list(csvReader1)

            if len(rows1) > 1:
                secondRow1 = rows1[1]
                modifiedRow1 = []
                for cell1 in secondRow1:
                    cellValue1 = str(cell1).strip()
                    for key1, value1 in sdw04Replacements.items():
                        if key1.lower() in cellValue1.lower():
                            cellValue1 = cellValue1.replace(key1, value1)
                    modifiedRow1.append(cellValue1)
                rows1[1] = modifiedRow1
            
        with open(newSDW04Template, 'w', newline="") as outputCSV1:
            csvWriter1 = csv.writer(outputCSV1)
            csvWriter1.writerows(rows1)

    except Exception as error:
        print(f"ERROR: {error}\n", traceback.format_exc())
        authLog.error(f"Error message: {error}\n", traceback.format_exc())
